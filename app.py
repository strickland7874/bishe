from pathlib import Path

from docx import Document
from docx.shared import Cm
from docx.oxml.ns import qn
import plotly.graph_objects as go
from dash_loading_spinners import Hash
from dash import Dash, dcc, html, dash_table, Input, Output, State

from tool import *

t = np.array([10, 20, 30, 40, 60, 80, 100, 120, 150, 210,
              270, 330, 400, 450, 645, 870, 990, 1185])
s = np.array([0.16, 0.48, 0.54, 0.65, 0.75, 1.00, 1.12, 1.22, 1.36, 1.55,
              1.70, 1.83, 1.89, 1.98, 2.17, 2.38, 2.46, 2.54])

external_stylesheets = ['https://codepen.io/chriddyp/pen/bWLwgP.css']
app = Dash(__name__, external_stylesheets=external_stylesheets)
app.title = '地下水动力学'
server = app.server
app.layout = html.Div([
    html.H1('求解水文地质参数'),

    # 下拉框
    dcc.Dropdown(
        id='dropdown',
        options=[
            {'label': '配线法', 'value': 'theis'},
            {'label': '直线图解法', 'value': 'line'},
        ],
        value='theis',
        clearable=False,
        searchable=False
    ),
    html.Br(),

    # 用户输入组件
    html.Label('抽水孔定流量Q（单位 m³/min ）'),
    dcc.Input(id='Q-input', value=1, type='number'),
    html.Br(),
    html.Br(),
    html.Label('输入观测孔于抽水孔距离r（单位 m ）'),
    dcc.Input(id='r-input', value=140, type='number'),
    html.Br(),
    html.Br(),

    # 计算结果
    html.Strong(id='result', style={'font-size': '22px'}),
    html.Br(),
    html.Br(),

    # 滚动条
    dcc.Markdown(id='slider1-output', mathjax=True),
    html.Div([dcc.Slider(id='slider1', min=-5, max=5, step=0.001, included=False,
                         marks={i: str(i) for i in range(-5, 5 + 1)}, tooltip={'placement': 'bottom'},
                         value=0.22)],
             style={'width': '20%'}),
    html.Br(),
    dcc.Markdown(id='slider-output-lgtr2', mathjax=True),
    html.Div([dcc.Slider(id='slider2', min=-5, max=5, step=0.001, included=False,
                         marks={i: str(i) for i in range(-5, 5 + 1)}, tooltip={'placement': 'bottom'},
                         value=3.31)],
             style={'width': '20%'}),
    html.Br(),

    # 图
    dcc.Graph(id='graph'),

    # 生成按钮
    html.Div(id='hidden', style={'display': 'none'}),
    html.Button(id='save', children='保存到Word', n_clicks=0),
    html.Br(),
    html.Br(),
    Hash(html.Label(id='spinner'), speed_multiplier=2),
    html.Br(),
    html.Br(),
    html.Br(),

    # 上传组件
    dcc.Upload(
        id='upload',
        children=html.Div([html.A('点击上传或将文件拖入此区域')]),
        style={
            'width': '100%', 'height': '60px', 'lineHeight': '60px',
            'borderWidth': '1px', 'borderStyle': 'dashed',
            'borderRadius': '5px', 'textAlign': 'center', 'margin': '10px'
        },
    ),

    # 数据表
    html.Div([dash_table.DataTable(id='datatable', editable=True,
                                   data=[{'t': i, 's': j} for i, j in zip(t, s)],
                                   columns=[{'name': 't', 'id': 't'}, {'name': 's', 'id': 's'}],
                                   style_cell={'fontSize': 20, 'font-family': 'sans-serif'})],
             style={'width': '20%'}),
])


@app.callback(
    Output('Q-input', 'value'),
    Output('r-input', 'value'),
    Output('slider1', 'value'),
    Output('slider2', 'value'),
    # Output('datatable', 'data'),
    Input('dropdown', 'value')
)
def initialize(fit_type):
    """选择不同解法时初始化数据"""
    if fit_type == 'theis':  # 配线法
        Q = 60 / 60
        r = 140
        # t = np.array([10, 20, 30, 40, 60, 80, 100, 120, 150, 210,
        #               270, 330, 400, 450, 645, 870, 990, 1185])
        # s = np.array([0.16, 0.48, 0.54, 0.65, 0.75, 1.00, 1.12, 1.22, 1.36, 1.55,
        #               1.70, 1.83, 1.89, 1.98, 2.17, 2.38, 2.46, 2.54])
        # data = [{'t': i, 's': j} for i, j in zip(t, s)]
        lgs = 0.22
        lgtr2 = 3.31
        return Q, r, lgs, lgtr2  # , data
    elif fit_type == 'line':  # 直线图解法
        Q = 528 / 1440
        r = 90
        # t = np.array([1, 2, 4, 6, 9, 20, 30, 40, 50, 60,
        #               90, 120, 150, 360, 550, 720])
        # s = np.array([0.025, 0.039, 0.061, 0.08, 0.0106, 0.168, 0.2, 0.226, 0.247,
        #               0.264, 0.304, 0.330, 0.350, 0.426, 0.440, 0.445])
        # data = [{'t': i, 's': j} for i, j in zip(t, s)]
        t0 = 3.0
        slope = 0.20426
        return Q, r, t0, slope  # , data


@app.callback(
    Output('slider1-output', 'children'),
    Input('dropdown', 'value'),
    Input('slider1', 'value')
)
def update_slider1(fit_type, value):
    """更新滚动条1"""
    if fit_type == 'theis':  # 配线法
        return r'$\Delta\lg s$：' + str(value)
    elif fit_type == 'line':  # 直线图解法
        return r'$t_0$：' + str(value)


@app.callback(
    Output('slider-output-lgtr2', 'children'),
    Input('dropdown', 'value'),
    Input('slider2', 'value')
)
def update_slider2(fit_type, value):
    """更新滚动条2"""
    if fit_type == 'theis':  # 配线法
        return r'$\Delta\lg\frac{t}{r^2}$：' + str(value)
    elif fit_type == 'line':  # 直线图解法
        return '$slope$：' + str(value)


@app.callback(
    Output('datatable', 'data'),
    Output('datatable', 'columns'),
    Input('upload', 'contents'),
    State('upload', 'filename'),
    prevent_initial_call=True
)
def update_datatable(contents, filename):
    """上传文件后更新表格"""
    if not contents:
        return [{}], []
    df = parse_content(contents, filename)
    return df.to_dict('records'), [{'name': i, 'id': i} for i in df.columns]


@app.callback(
    Output('result', 'children'),
    Input('dropdown', 'value'),
    Input('Q-input', 'value'),
    Input('r-input', 'value'),
    Input('slider1', 'value'),
    Input('slider2', 'value')
)
def update_result(fit_type, Q, r, slider1, slider2):
    """计算导水系数和储水系数"""
    T, S = 0, 0
    if fit_type == 'theis':  # 配线法
        lgs = slider1
        lgtr2 = slider2
        T, S = theis_calculate(Q, lgs, lgtr2)
    elif fit_type == 'line':  # 直线图解法
        t0 = slider1
        slope = slider2
        T, S = line_calculate(Q, r, t0, slope)
    return f'导水系数 T = {T:.4f} m²/min\n储水系数 S = {S:.4e}'


@app.callback(
    Output('graph', 'figure'),
    Input('dropdown', 'value'),
    Input('datatable', 'data'),
    Input('Q-input', 'value'),
    Input('r-input', 'value'),
    Input('slider1', 'value'),
    Input('slider2', 'value')
)
def update_figure(fit_type, data, Q, r, slider1, slider2):
    """画图"""
    fig = go.Figure()

    t = np.array([i['t'] for i in data])
    s = np.array([i['s'] for i in data])

    if fit_type == 'theis':  # 配线法
        xmin = 1.0e-1
        xmax = 1.0e4
        ymin = 1.0e-2
        ymax = 10.0
        lgs = slider1
        lgtr2 = slider2
        line_x, line_y, scatter_x, scatter_y = theis_fit(t, s, Q, r, lgs, lgtr2)
        fig = go.Figure()
        fig.add_trace(go.Scatter(x=line_x, y=line_y, mode='lines', name='Theis'))
        fig.add_trace(go.Scatter(x=scatter_x, y=scatter_y, mode='markers', name='observed',
                                 marker=dict(symbol='star-dot', size=12)))
        fig.update_layout(xaxis_range=[np.log10(xmin), np.log10(xmax)])
        fig.update_layout(yaxis_range=[np.log10(ymin), np.log10(ymax)])
        fig.update_xaxes(type='log')
        fig.update_yaxes(type='log')

    elif fit_type == 'line':  # 直线图解法
        t0 = slider1  # 3.0
        slope = slider2  # 0.20426
        line_x, line_y, scatter_x, scatter_y = line_fit(t, s, t0, slope)
        fig = go.Figure()
        fig.add_trace(go.Scatter(x=line_x, y=line_y, mode='lines', name='Jacob fit'))
        fig.add_trace(go.Scatter(x=scatter_x, y=scatter_y, mode='markers', name='observed',
                                 marker=dict(symbol='star-dot', size=12)))
        fig.update_xaxes(type='log')

    return fig


@app.callback(
    Output('spinner', 'children'),
    State('dropdown', 'value'),
    Input('save', 'n_clicks'),
    State('graph', 'figure'),
    State('Q-input', 'value'),
    State('r-input', 'value'),
    State('slider1', 'value'),
    State('slider2', 'value'),
    prevent_initial_call=True
)
def save_to_word(fit_type, n_clicks, figure, Q, r, slider1, slider2):
    folder = Path('output')
    folder.mkdir(parents=True, exist_ok=True)
    filepath, imagepath = None, None
    fig = go.Figure(**figure)
    T, S = 0, 0
    if fit_type == 'theis':  # 配线法
        imagepath = str(folder / '配线法.jpg')
        filepath = str(folder / '配线法.docx')
        lgs = slider1
        lgtr2 = slider2
        T, S = theis_calculate(Q, lgs, lgtr2)
    elif fit_type == 'line':  # 直线图解法
        imagepath = str(folder / '直线图解法.jpg')
        filepath = str(folder / '直线图解法.docx')
        t0 = slider1
        slope = slider2
        T, S = line_calculate(Q, r, t0, slope)

    fig.write_image(imagepath, scale=2)  # 保存图片
    document = Document()
    style = document.styles['Normal']
    style.font.name = '宋体'
    style.font.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    document.add_paragraph('项目名称：')
    document.add_paragraph('抽水试验时间：')
    document.add_paragraph('利用抽水试验实测曲线与标准曲线进行拟合，可求出相应水文地质参数，拟合结果如下图所示：')
    document.add_picture(imagepath, width=Cm(15.3))  # 插入图片
    document.add_paragraph('水文地质参数求解结果如下表所示')
    table = document.add_table(rows=2, cols=7, style='Table Grid')
    texts = ['含水层位', '观测井', '抽水孔与观测孔间距r', '抽水孔流量Q', '含水层厚度', '导水系数T', '储水系数S']
    for cell, text in zip(table.rows[0].cells, texts):
        cell.text = text
    texts = ['', '', f'{r}m', f'{Q:.4f}m³/min', '', f'{T:.4f}m³/min', f'{S:.4e}']
    for cell, text in zip(table.rows[1].cells, texts):
        cell.text = text
    document.save(filepath)  # 保存

    return f'导出成功：{filepath}'


if __name__ == '__main__':
    app.run_server(debug=True)
